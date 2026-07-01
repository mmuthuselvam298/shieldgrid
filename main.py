import os
import re
import uuid
import tempfile
from datetime import datetime
from xml.sax.saxutils import escape
import pdfplumber
from docx import Document
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse
from pydantic import BaseModel
from paddleocr import PaddleOCR
from presidio_analyzer import AnalyzerEngine, Pattern, PatternRecognizer
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_anonymizer import AnonymizerEngine
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from fastapi.staticfiles import StaticFiles

# =====================================================================
# ✅ DEFERRED ENGINES SETUP (Lazy Loading to bypass boot timeouts)
# =====================================================================
ocr = None
analyzer = None
anonymizer = AnonymizerEngine()  # Lightweight engine; safe to remain global

# ✅ Staging array to collect custom patterns without triggering spaCy instantiation
recognizers = []

def extract_text_from_file(filename, content):
    global ocr  # Access the deferred global variable
    filename = filename.lower()
    temp_suffix = uuid.uuid4().hex
    temp_dir = tempfile.gettempdir()
    temp_docx = os.path.join(temp_dir, f"temp_{temp_suffix}.docx")
    temp_pdf = os.path.join(temp_dir, f"temp_{temp_suffix}.pdf")
    temp_image = os.path.join(temp_dir, f"temp_{temp_suffix}.jpg")

    try:
        if filename.endswith(".txt"):
            return content.decode("utf-8", errors="ignore")

        if filename.endswith(".docx"):
            with open(temp_docx, "wb") as f:
                f.write(content)

            doc = Document(temp_docx)
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            return text

        if filename.endswith(".pdf"):
            with open(temp_pdf, "wb") as f:
                f.write(content)

            text = ""
            with pdfplumber.open(temp_pdf) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        page_text = re.sub(r"\s+", " ", page_text)
                        text += page_text + "\n"
            return text

        if filename.endswith((".jpg", ".jpeg", ".png")):
            with open(temp_image, "wb") as f:
                f.write(content)

            if ocr is None:
                ocr = PaddleOCR(
                    use_angle_cls=True,
                    lang="en"
                )

            result = ocr.ocr(temp_image)
            text = ""
            if result:
                page = result[0]
                if isinstance(page, dict):
                    text = "\n".join(page.get("rec_texts", []))
            return text

        return ""
    finally:
        for temp_file in [temp_pdf, temp_docx, temp_image]:
            if os.path.exists(temp_file):
                try:
                    os.remove(temp_file)
                except Exception:
                    pass

app = FastAPI(title="ShieldGrid PII Detection & Redaction System")

# Mount static files (CSS, JS, assets)
app.mount("/static", StaticFiles(directory="static"), name="static")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def add_recognizer(entity, regex, score=0.9):
    pattern = Pattern(name=entity.lower(), regex=regex, score=score)
    recognizer = PatternRecognizer(supported_entity=entity, patterns=[pattern])
    recognizers.append(recognizer)

GENERAL_ENTITIES = {"PERSON", "PHONE_NUMBER", "EMAIL_ADDRESS", "LOCATION", "PASSPORT_NUMBER"}

FINANCIAL_ENTITIES = {
    "PAN_NUMBER", "AADHAAR_NUMBER", "IFSC_CODE", "CARD_NUMBER", "TRANSACTION_ID",
    "BANK_ACCOUNT", "CUSTOMER_ID", "INSURANCE_POLICY", "CVV", "CHEQUE_NUMBER",
    "LOAN_ACCOUNT", "TAX_ID", "CUSTOMER_NAME", "SWIFT_CODE", "CARD_EXPIRY",
    "GST_NUMBER", "MICR_CODE", "LOAN_REFERENCE", "FD_NUMBER", "REFERENCE_NUMBER",
    "UPI_ID", "DEMAT_ACCOUNT", "IBAN", "WALLET_ID", "MERCHANT_ID", "BRANCH_CODE",
    "DATE_OF_BIRTH", "PASSPORT_NUMBER",
}

HEALTHCARE_ENTITIES = {
    "MEDICAL_RECORD_NUMBER", "HEALTH_ID", "PATIENT_ID", "PATIENT_NAME", "DOCTOR_NAME",
    "DOCTOR_LICENSE", "HOSPITAL_ID", "HOSPITAL_NAME", "PRESCRIPTION_ID", "LAB_REPORT_ID",
    "HEALTH_INSURANCE", "DIAGNOSIS_CODE", "VACCINE_ID", "CLINICAL_TRIAL", "BED_NUMBER",
    "BLOOD_GROUP", "DATE_OF_BIRTH", "PHONE_NUMBER", "EMAIL_ADDRESS", "PERSON",
    "LOCATION", "PASSPORT_NUMBER", "WARD_NUMBER", "ROOM_NUMBER", "OP_NUMBER",
    "IP_NUMBER", "UHID", "CLAIM_ID", "DISCHARGE_ID", "INSURANCE_MEMBER_ID", "ADMISSION_NUMBER",
    "AGE", "GENDER", "HEIGHT", "WEIGHT", "INSURANCE_PROVIDER", "INSURANCE_POLICY", "RELATIONSHIP"
}

# --- Financial Recognizers ---
add_recognizer("PAN_NUMBER", r"\b[A-Z]{5}[0-9]{4}[A-Z]\b")
add_recognizer("AADHAAR_NUMBER", r"\b\d{4}[- ]?\d{4}[- ]?\d{4}\b")
add_recognizer("IFSC_CODE", r"\b[A-Z]{4}0[A-Z0-9]{6}\b")
add_recognizer("PASSPORT_NUMBER", r"\b[A-Z][0-9]{7}\b")
add_recognizer("CARD_NUMBER", r"\b(?:\d{4}[- ]?){3}\d{4}\b")
add_recognizer("BANK_ACCOUNT", r"\b\d{11,18}\b")
add_recognizer("CUSTOMER_ID", r"\bCUST\d{4,12}\b")
add_recognizer("CVV", r"\b\d{3}\b")
add_recognizer("CHEQUE_NUMBER", r"\b\d{6}\b")
add_recognizer("LOAN_ACCOUNT", r"\bLOAN\d{6,15}\b")
add_recognizer("TAX_ID", r"TAX\d+")
add_recognizer("EMPLOYEE_ID", r"EMP\d+")
add_recognizer("SWIFT_CODE", r"\b(?!CUSTOMER\b|PASSPORT\b)[A-Z]{4}[A-Z]{2}[A-Z0-9]{2}([A-Z0-9]{3})?\b", score=0.6)
add_recognizer("CARD_EXPIRY", r"\b(0[1-9]|1[0-2])/([0-9]{2})\b")
add_recognizer("GST_NUMBER", r"\b\d{2}[A-Z]{5}\d{4}[A-Z][A-Z\d]Z[A-Z\d]\b")
add_recognizer("MICR_CODE", r"\b\d{9}\b")
add_recognizer("FD_NUMBER", r"\bFD\d{6,15}\b")
add_recognizer("REFERENCE_NUMBER", r"\bREF\d+\b")
add_recognizer("UPI_ID", r"\b[a-zA-Z0-9._-]{2,}@(upi|ybl|ibl|oksbi|okhdfcbank|okicici|paytm|axl)\b")
add_recognizer("DEMAT_ACCOUNT", r"\bIN\d{14}\b")
add_recognizer("IBAN", r"\b[A-Z]{2}\d{2}[A-Z0-9]{11,30}\b")
add_recognizer("WALLET_ID", r"\bWALLET\d{5,15}\b")
add_recognizer("MERCHANT_ID", r"\bMID\d{6,15}\b")
add_recognizer("BRANCH_CODE", r"\bBR\d{3,8}\b")
add_recognizer("DATE_OF_BIRTH", r"\d{2}/\d{2}/\d{4}")
add_recognizer("TRANSACTION_ID", r"\b(?:TXN|TRX|UTR|NEFT|RTGS|IMPS)[A-Z0-9]{6,20}\b")
add_recognizer("LOAN_REFERENCE", r"\b(?:LN|LOAN)\d{6,15}\b")
add_recognizer("CUSTOMER_NAME", r"(?<=Customer Name:\s)[A-Z][a-zA-Z'.]*(?:\s[A-Z][a-zA-Z'.]*)*", score=0.95)

# --- General Person Recognizers ---
add_recognizer("PERSON", r"(?<=my name is\s)[A-Z][a-zA-Z]+(?:\s[A-Z][a-zA-Z]+)*", score=0.95)
add_recognizer("PERSON", r"(?<=Name:\s)[A-Z][a-zA-Z]+(?:\s[A-Z][a-zA-Z]+)*", score=0.95)

# --- Healthcare Recognizers ---
add_recognizer("MEDICAL_RECORD_NUMBER", r"\bMRN-\d{5,10}\b")
add_recognizer("HEALTH_ID", r"\bHID-\d{4}-\d{4}-\d{4}\b")
add_recognizer("PATIENT_ID", r"(?<=Patient ID:\s)[A-Z0-9-]+", score=0.95)
add_recognizer("HOSPITAL_ID", r"\bHOSP\d{5,10}\b")
add_recognizer("DOCTOR_LICENSE", r"\bDOC\d{5,10}\b")
add_recognizer("PRESCRIPTION_ID", r"\bRX\d{6,12}\b")
add_recognizer("LAB_REPORT_ID", r"\bLAB\d{5,12}\b")
add_recognizer("HEALTH_INSURANCE", r"\bHI\d{8,15}\b")
add_recognizer("VACCINE_ID", r"\bVAC\d{6,12}\b")
add_recognizer("CLINICAL_TRIAL", r"\bCT-\d{4,10}\b")
add_recognizer("BED_NUMBER", r"\bBED-\d{1,4}\b")
add_recognizer("OP_NUMBER", r"\bOP\d{4,10}\b", score=0.9)
add_recognizer("IP_NUMBER", r"\bIP\d{4,10}\b", score=0.9)
add_recognizer("UHID", r"\bUHID[- ]?\d{5,15}\b", score=0.9)
add_recognizer("CLAIM_ID", r"\bCLAIM[- ]?\d{5,15}\b", score=0.9)
add_recognizer("DISCHARGE_ID", r"\bDIS\d{4,10}\b", score=0.9)
add_recognizer("INSURANCE_MEMBER_ID", r"\bMEM\d{5,15}\b", score=0.9)
add_recognizer("ADMISSION_NUMBER", r"\bADM\d{4,10}\b", score=0.9)
add_recognizer("HOSPITAL_NAME", r"(?<=Hospital:\s)[^\n]+", score=0.95)
add_recognizer("DIAGNOSIS_CODE", r"\b[A-Z][0-9]{2}(?:\.[0-9A-Z]{1,4})?\b")
add_recognizer("WARD_NUMBER", r"\bWard(?:\s*No\.?)?[- ]?\d+\b", score=0.9)
add_recognizer("ROOM_NUMBER", r"\bRoom(?:\s*No\.?)?[- ]?\d+\b", score=0.9)
add_recognizer("BLOOD_GROUP", r"\b(?:A|B|AB|O)[+-]\b")
add_recognizer("PATIENT_NAME", r"(?<=Patient Name:\s)[A-Z][a-zA-Z'.]*(?:\s[A-Z][a-zA-Z'.]*)*", score=0.95)
add_recognizer("DOCTOR_NAME", r"(?<=(?:Doctor|Consultant):\s)(?:Dr\.\s)?[A-Z][a-zA-Z'.]*(?:\s[A-Z][a-zA-Z'.]*)*", score=0.95)

# New demographic & insurance recognizers to catch remaining text indicators
add_recognizer("AGE", r"\bAge:\s*\d+(?:\s*years?)?\b", score=0.95)
add_recognizer("GENDER", r"\bGender:\s*(?:Male|Female|Other)\b", score=0.95)
add_recognizer("HEIGHT", r"\bHeight:\s*\d+\s*cm\b", score=0.9)
add_recognizer("WEIGHT", r"\bWeight:\s*\d+\s*kg\b", score=0.9)
add_recognizer("INSURANCE_PROVIDER", r"(?<=Health Insurance Provider:\s)[^\n]+", score=0.95)
add_recognizer("INSURANCE_POLICY", r"(?<=Policy Number:\s)[A-Z0-9]+", score=0.95)
add_recognizer("RELATIONSHIP", r"\bRelationship:\s*[A-Za-z]+\b", score=0.85)

COMPLIANCE_MAP = {
    "finance": FINANCIAL_ENTITIES | GENERAL_ENTITIES,
    "healthcare": HEALTHCARE_ENTITIES | GENERAL_ENTITIES,
    "general": GENERAL_ENTITIES,
}

class TextRequest(BaseModel):
    text: str
    compliance: str = "general"

class DownloadRequest(BaseModel):
    text: str
    format: str
    compliance: str
    entities_count: int = 0
    entity_summary: dict = {}

def redact_text(text, compliance):
    global analyzer  # Access the deferred global variable
    compliance = compliance.lower()
    address_block_pattern = (
        r"(?is)"
        r"([A-Za-z ]*Address:\s*)"
        r"(.*?)"
        r"(?=\n[A-Za-z][A-Za-z ]{2,30}:\s|\Z)"
    )

    address_matches = list(re.finditer(address_block_pattern, text))
    address_segments = [match.group(2) for match in address_matches]

    text_for_presidio = re.sub(address_block_pattern, r"\1<ADDRESS>", text)

    text_for_presidio = re.sub(
        r"(?i)(Customer Name:\s*[^\n]+?)\s*(?:[:\-]\s*)?(CUST\d+)",
        r"\1\nCustomer ID: \2",
        text_for_presidio
    )

    if analyzer is None:
        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}]
        }
        provider = NlpEngineProvider(nlp_configuration=configuration)
        nlp_engine = provider.create_engine()
        analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
        
        for recognizer in recognizers:
            analyzer.registry.add_recognizer(recognizer)

    results = analyzer.analyze(text=text_for_presidio, language="en")
    allowed_entities = COMPLIANCE_MAP.get(compliance, GENERAL_ENTITIES)

    filtered_results = [
        r
        for r in results
        if r.entity_type not in {"US_DRIVER_LICENSE", "UK_NHS"} and r.entity_type in allowed_entities
    ]

    entity_summary = {}
    for r in filtered_results:
        entity_summary[r.entity_type] = entity_summary.get(r.entity_type, 0) + 1

    if address_segments:
        entity_summary["ADDRESS"] = entity_summary.get("ADDRESS", 0) + len(address_segments)

    seen = set()
    detected_entities = []
    for r in filtered_results:
        key = (r.entity_type, r.start, r.end)
        if key not in seen:
            seen.add(key)
            detected_entities.append(
                {
                    "entity": r.entity_type,
                    "score": round(r.score, 2),
                    "fragment": text_for_presidio[r.start : r.end],
                }
            )

    for chunk in address_segments:
        detected_entities.append(
            {
                "entity": "ADDRESS",
                "score": 1.0,
                "fragment": chunk.strip()
            }
        )

    detected_entities.sort(key=lambda x: x["entity"])
    anonymized = anonymizer.anonymize(text=text_for_presidio, analyzer_results=filtered_results)

    return {
        "original_text": text,
        "anonymized_text": anonymized.text,
        "detected_entities": detected_entities,
        "entity_summary": entity_summary,
    }

def remove_generated_file(path: str):
    if os.path.exists(path):
        try:
            os.remove(path)
        except Exception:
            pass

# --- Primary Application Gateway Endpoints ---

@app.get("/", response_class=HTMLResponse)
def home():
    """Serves the front-end user interface dashboard."""
    with open("templates/index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.get("/health")
def health_check():
    """Isolated backend heartbeat verification node."""
    return {"status": "ready", "service": "presidio-redaction"}

@app.post("/analyze")
def analyze_text(request: TextRequest):
    return redact_text(request.text, request.compliance)

@app.post("/analyze-file")
async def analyze_file(file: UploadFile = File(...), compliance: str = Form("general")):
    content = await file.read()
    text = extract_text_from_file(file.filename, content)
    return redact_text(text, compliance)

@app.post("/download")
def download_redacted_file(request: DownloadRequest, background_tasks: BackgroundTasks):
    fmt = request.format.lower()
    text_content = request.text
    compliance_slug = request.compliance.lower()
    compliance_mode = request.compliance.upper()
    ent_count = request.entities_count
    summary_data = request.entity_summary
    now_ts = datetime.now()
    date_stamp = now_ts.strftime("%Y-%m-%d")
    time_stamp = now_ts.strftime("%H%M%S")
    current_time_str = now_ts.strftime("%d %B %Y, %H:%M:%S")

    filename = f"{compliance_slug}_redacted_{date_stamp}_{time_stamp}_{uuid.uuid4().hex[:8]}.{fmt}"

    temp_dir = tempfile.gettempdir()
    filepath = os.path.join(temp_dir, filename)

    ledger_text_lines = "ENTITY SUMMARY\n"
    if summary_data:
        for ent, count in summary_data.items():
            label = ent.replace("_", " ").title()
            ledger_text_lines += f"{label.ljust(25, '.')} {str(count).rjust(3, ' ')}\n"
    else:
        ledger_text_lines += "No sensitive entities isolated.\n"

    metadata_header = (
        "==================================================\n"
        f"{compliance_mode} COMPLIANCE REPORT\n"
        "GENERATED BY               : ShieldGrid Privacy Engine\n"
        f"GENERATION DATE            : {current_time_str}\n"
        f"SENSITIVE ENTITIES DETECTED: {ent_count}\n"
        "==================================================\n\n"
        f"{ledger_text_lines}"
        "==================================================\n\n"
    )

    if fmt == "txt":
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(metadata_header + text_content)
        media = "text/plain"

    elif fmt == "docx":
        doc = Document()
        doc.add_heading(f"{compliance_mode} Compliance Report", level=1)
        doc.add_paragraph(f"Compliance Profile: {compliance_mode}")
        doc.add_paragraph("Generated By: ShieldGrid Privacy Engine")
        doc.add_paragraph(f"Generation Date: {current_time_str}")
        doc.add_paragraph(f"Sensitive Entities Detected: {ent_count}")
        doc.add_paragraph("-" * 40)
        doc.add_heading("Entity Summary Ledger", level=2)
        if summary_data:
            for ent, count in summary_data.items():
                doc.add_paragraph(f"- {ent.replace('_', ' ').title()}: {count}")
        else:
            doc.add_paragraph("No sensitive metrics isolated.")
        doc.add_paragraph("-" * 40)
        doc.add_paragraph("")
        for line in text_content.splitlines():
            doc.add_paragraph(line)
        doc.save(filepath)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    elif fmt == "pdf":
        doc = SimpleDocTemplate(filepath)
        styles = getSampleStyleSheet()
        story = [
            Paragraph(f"<b>{compliance_mode} Compliance Report</b>", styles["Heading1"]),
            Paragraph(f"<b>Compliance Profile:</b> {compliance_mode}", styles["Normal"]),
            Paragraph("<b>Generated By:</b> ShieldGrid Privacy Engine", styles["Normal"]),
            Paragraph(f"<b>Generation Date:</b> {current_time_str}", styles["Normal"]),
            Paragraph(f"<b>Sensitive Entities Detected:</b> {ent_count}", styles["Normal"]),
            Spacer(1, 10),
            Paragraph("<b>Entity Summary Ledger:</b>", styles["Heading2"]),
        ]
        if summary_data:
            for ent, count in summary_data.items():
                label = ent.replace("_", " ").title()
                story.append(Paragraph(f"&bull; {label}: {count}", styles["Normal"]))
        else:
            story.append(Paragraph("No sensitive metrics isolated.", styles["Normal"]))

        story.extend([Spacer(1, 10), Paragraph("<hr/>", styles["Normal"]), Spacer(1, 15)])
        for line in text_content.splitlines():
            clean_line = escape(line).replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;").replace(" ", "&nbsp;")
            if not clean_line.strip():
                story.append(Spacer(1, 10))
            else:
                story.append(Paragraph(clean_line, styles["BodyText"]))
        doc.build(story)
        media = "application/pdf"

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format profile type: {request.format}")

    background_tasks.add_task(remove_generated_file, filepath)
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type=media,
        headers={
            "Content-Disposition": f'attachment; filename="{filename}"',
            "Access-Control-Expose-Headers": "Content-Disposition",
        },
    )